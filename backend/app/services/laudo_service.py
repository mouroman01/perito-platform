from io import BytesIO

from docx import Document
from fastapi import HTTPException, status
from fpdf import FPDF
from sqlalchemy import select
from sqlalchemy.orm import Session, joinedload

from app.models.laudo import Laudo
from app.models.modelo import Modelo
from app.models.usuario import Usuario
from app.schemas.laudo import LaudoCreate, LaudoUpdate


def _query_base():
    return select(Laudo).options(
        joinedload(Laudo.processo),
        joinedload(Laudo.modelo),
        joinedload(Laudo.criado_por),
    )


def listar(db: Session, processo_id: int | None = None) -> list[Laudo]:
    query = _query_base().order_by(Laudo.criado_em.desc())
    if processo_id is not None:
        query = query.where(Laudo.processo_id == processo_id)
    return list(db.scalars(query))


def obter_ou_404(db: Session, laudo_id: int) -> Laudo:
    laudo = db.scalar(_query_base().where(Laudo.id == laudo_id))
    if laudo is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Laudo não encontrado")
    return laudo


def criar(db: Session, dados: LaudoCreate, usuario: Usuario) -> Laudo:
    conteudo = dados.conteudo
    if not conteudo and dados.modelo_id is not None:
        modelo = db.get(Modelo, dados.modelo_id)
        if modelo is not None:
            conteudo = modelo.conteudo

    laudo = Laudo(
        titulo=dados.titulo,
        processo_id=dados.processo_id,
        modelo_id=dados.modelo_id,
        conteudo=conteudo,
        criado_por_id=usuario.id,
    )
    db.add(laudo)
    db.commit()
    db.refresh(laudo)
    return obter_ou_404(db, laudo.id)


def atualizar(db: Session, laudo_id: int, dados: LaudoUpdate) -> Laudo:
    laudo = obter_ou_404(db, laudo_id)
    for campo, valor in dados.model_dump(exclude_unset=True).items():
        setattr(laudo, campo, valor)
    db.commit()
    return obter_ou_404(db, laudo_id)


def remover(db: Session, laudo_id: int) -> None:
    laudo = obter_ou_404(db, laudo_id)
    db.delete(laudo)
    db.commit()


def gerar_pdf(laudo: Laudo) -> bytes:
    pdf = FPDF()
    pdf.set_margin(20)
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 16)
    pdf.multi_cell(0, 10, laudo.titulo)
    pdf.ln(2)

    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(90, 90, 90)
    pdf.cell(0, 6, f"Processo: {laudo.processo.numero}", new_x="LMARGIN", new_y="NEXT")
    if laudo.criado_por:
        pdf.cell(0, 6, f"Elaborado por: {laudo.criado_por.nome}", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Helvetica", "", 11)
    conteudo = laudo.conteudo.encode("latin-1", "replace").decode("latin-1")
    pdf.multi_cell(0, 6, conteudo)

    return bytes(pdf.output())


def gerar_docx(laudo: Laudo) -> bytes:
    documento = Document()
    documento.add_heading(laudo.titulo, level=1)

    p_processo = documento.add_paragraph()
    p_processo.add_run(f"Processo: {laudo.processo.numero}").italic = True

    if laudo.criado_por:
        p_autor = documento.add_paragraph()
        p_autor.add_run(f"Elaborado por: {laudo.criado_por.nome}").italic = True

    documento.add_paragraph()

    for linha in laudo.conteudo.split("\n"):
        documento.add_paragraph(linha)

    buffer = BytesIO()
    documento.save(buffer)
    return buffer.getvalue()
